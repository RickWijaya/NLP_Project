"""
Text extraction module for different file types.
Supports PDF, DOCX, TXT, and XLSX files.
"""

import os
import re
from pathlib import Path
from typing import Optional, List

import pdfplumber
from docx import Document
from openpyxl import load_workbook

from app.utils.logger import logger, log_processing_step


class TextExtractor:
    """Extracts text content from various document formats."""
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".xlsx"}
    
    def __init__(self):
        """Initialize the text extractor."""
        pass
    
    def extract(self, file_path: str, document_id: str = "") -> str:
        """
        Extract text from a file.
        
        Args:
            file_path: Path to the file
            document_id: Optional document ID for logging
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")
        
        log_processing_step(document_id, "extraction", "started", f"Extracting from {extension} file")
        
        try:
            if extension == ".pdf":
                segments = self._extract_pdf(file_path)
            elif extension == ".docx":
                segments = [{"content": self._extract_docx(file_path), "page_label": "1"}]
            elif extension == ".txt":
                segments = [{"content": self._extract_txt(file_path), "page_label": "1"}]
            elif extension == ".xlsx":
                segments = [{"content": self._extract_xlsx(file_path), "page_label": "1"}]
            else:
                raise ValueError(f"Unsupported file type: {extension}")
            
            # Clean up the text in each segment
            for seg in segments:
                seg["content"] = self._clean_text(seg["content"])
            
            # Filter out empty segments
            segments = [s for s in segments if s["content"].strip()]
            
            total_chars = sum(len(s["content"]) for s in segments)
            log_processing_step(
                document_id, 
                "extraction", 
                "completed", 
                f"Extracted {total_chars} characters across {len(segments)} segments"
            )
            
            return segments
            
        except Exception as e:
            log_processing_step(document_id, "extraction", "failed", str(e))
            raise
    
    def _extract_pdf(self, file_path: str) -> List[dict]:
        """
        Extract text from PDF file using pdfplumber, preserved per page.
        """
        segments = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract text with better layout handling
                    page_text = page.extract_text(
                        layout=True,
                        x_tolerance=3,
                        y_tolerance=3
                    )
                    
                    if page_text:
                        # Clean up layout artifacts
                        lines = page_text.split('\n')
                        cleaned_lines = [re.sub(r'\s{3,}', '  ', line.strip()) for line in lines if line.strip()]
                        content = '\n'.join(cleaned_lines)
                        
                        if content:
                            segments.append({
                                "content": content,
                                "page_label": str(page_num + 1)
                            })
                            
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying fallback: {e}")
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            segments.append({
                                "content": page_text,
                                "page_label": str(page_num + 1)
                            })
            except Exception as e2:
                logger.error(f"PDF extraction failed: {e2}")
                raise
        
        return segments
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() 
                    for cell in row.cells 
                    if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    def _extract_txt(self, file_path: str) -> str:
        """
        Extract text from TXT file with improved encoding detection.
        """
        # Try common encodings in order of likelihood
        encodings = ["utf-8", "utf-8-sig", "utf-16", "latin-1", "cp1252", "iso-8859-1"]
        
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    content = f.read()
                    # Verify it's valid text (no null bytes, printable)
                    if '\x00' not in content:
                        logger.info(f"TXT extracted with encoding: {encoding}")
                        return content
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        # Fallback with error handling
        logger.warning("Using fallback encoding with error handling")
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    
    def _extract_xlsx(self, file_path: str) -> str:
        """Extract text from XLSX file (textual content only)."""
        workbook = load_workbook(file_path, data_only=True)
        text_parts = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"Sheet: {sheet_name}")
            
            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value).strip())
                
                if row_values:
                    text_parts.append(" | ".join(row_values))
        
        return "\n".join(text_parts)
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text for better processing."""
        if not text:
            return ""
        
        # Replace multiple newlines with double newline (paragraph break)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Replace tabs with spaces
        text = text.replace('\t', ' ')
        
        # Replace multiple spaces with single space (preserve newlines)
        text = re.sub(r'[^\S\n]+', ' ', text)
        
        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # Remove empty lines at start and end
        text = text.strip()
        
        return text


# Global instance
text_extractor = TextExtractor()

